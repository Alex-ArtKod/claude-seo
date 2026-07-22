#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
md2docx.py — convert a Markdown article into a .docx file using python-docx.

Designed for the /seo-writer pipeline. Runs in a separate session/agent so it
does not consume the writer's context. No external binaries required
(pandoc / LibreOffice are NOT used).

Supported Markdown:
- # / ## / ### / #### headings        -> Word heading styles
- paragraphs                          -> normal paragraphs
- pipe tables (| a | b |)             -> Word tables with header row
- unordered lists (-, *, +)           -> "List Bullet"
- ordered lists (1. 2. ...)           -> "List Number"
- inline **bold**, *italic*, `code`   -> runs with formatting
- <!-- comments --> and HTML lines    -> skipped

Usage:
    python md2docx.py input.md [output.docx]

If output is omitted, the input name with .docx extension is used.
Exit code 0 on success, 1 on error.
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    sys.stderr.write(
        "python-docx is not installed. Run: pip install python-docx\n"
    )
    sys.exit(1)


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
UL_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
OL_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")
COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)
INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def split_table_row(line):
    """Split a markdown table row into trimmed cells."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def add_inline_runs(paragraph, text):
    """Add text to a paragraph, honoring **bold**, *italic*, `code`."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:  # *italic*
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def flush_table(doc, rows):
    """Render collected markdown table rows into a Word table."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncols):
            text = row[j] if j < len(row) else ""
            para = cells[j].paragraphs[0]
            add_inline_runs(para, text)
            if i == 0:
                for run in para.runs:
                    run.bold = True


def convert(md_text, doc):
    lines = COMMENT_RE.sub("", md_text).split("\n")
    table_buffer = []
    i = 0
    n = len(lines)

    while i < n:
        raw = lines[i]
        line = raw.rstrip()

        # Inside / start of a table block
        if TABLE_ROW_RE.match(line):
            cells = split_table_row(line)
            # skip the |---|---| separator row
            if not TABLE_SEP_RE.match(line):
                table_buffer.append(cells)
            i += 1
            continue
        elif table_buffer:
            flush_table(doc, table_buffer)
            table_buffer = []

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Skip raw HTML lines
        if stripped.startswith("<") and stripped.endswith(">"):
            i += 1
            continue

        m = HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=min(level, 4))
            i += 1
            continue

        m = UL_RE.match(line)
        if m:
            para = doc.add_paragraph(style="List Bullet")
            add_inline_runs(para, m.group(1).strip())
            i += 1
            continue

        m = OL_RE.match(line)
        if m:
            para = doc.add_paragraph(style="List Number")
            add_inline_runs(para, m.group(1).strip())
            i += 1
            continue

        # Plain paragraph
        para = doc.add_paragraph()
        add_inline_runs(para, stripped)
        i += 1

    if table_buffer:
        flush_table(doc, table_buffer)


def main(argv):
    if len(argv) < 2:
        sys.stderr.write("Usage: python md2docx.py input.md [output.docx]\n")
        return 1

    in_path = Path(argv[1])
    if not in_path.is_file():
        sys.stderr.write(f"Input file not found: {in_path}\n")
        return 1

    out_path = Path(argv[2]) if len(argv) > 2 else in_path.with_suffix(".docx")

    md_text = in_path.read_text(encoding="utf-8")
    doc = Document()
    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    convert(md_text, doc)
    doc.save(str(out_path))
    print(f"OK: {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
